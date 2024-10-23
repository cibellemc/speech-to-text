import time
import streamlit as st
import docx
from sqlalchemy import text
from datetime import datetime
import io
import matplotlib.colors as mcolors
from pages.select import display_transcriptions, fetch_transcription_by_file_name
from services.database import conn
from transcriber import transcribe


def _style_language_uploader():
    languages = {
        "PT-BR": {
            "button": "Selecionar arquivos",
            "instructions": "Arraste e solte os arquivos aqui",
            "limits": "Limite de 1GB por arquivo | MP4, M4A, MP3, WAV",
        },
    }

    hide_label = (
        """
        <style>
            div[data-testid="stFileUploader"]>section[data-testid="stFileUploaderDropzone"]>button[data-testid="baseButton-secondary"] {
               color:white;
            }
            # div[data-testid="stFileUploader"]>section[data-testid="stFileUploaderDropzone"]>button[data-testid="baseButton-secondary"]::after {
            #     content: "BUTTON_TEXT";
            #     color:black;
            #     display: block;
            #     position: absolute;
            # }
            div[data-testid="stFileUploaderDropzoneInstructions"]>div>span {
               visibility:hidden;
            }
            div[data-testid="stFileUploaderDropzoneInstructions"]>div>span::after {
               content:"INSTRUCTIONS_TEXT";
               visibility:visible;
               display:block;
            }
             div[data-testid="stFileUploaderDropzoneInstructions"]>div>small {
               visibility:hidden;
            }
            div[data-testid="stFileUploaderDropzoneInstructions"]>div>small::before {
               content:"FILE_LIMITS";
               visibility:visible;
               display:block;
            }
        </style>
        """.replace(
            "BUTTON_TEXT", languages.get("PT-BR").get("button")
        )
        .replace("INSTRUCTIONS_TEXT", languages.get("PT-BR").get("instructions"))
        .replace("FILE_LIMITS", languages.get("PT-BR").get("limits"))
    )

    st.markdown(hide_label, unsafe_allow_html=True)


def group_speaker_segments(segments):
    grouped_transcription = []
    current_speaker = None
    current_text = ""

    for segment in segments:
        speaker = segment["speaker"]
        text = segment["text"]

        # Se o falante atual é o mesmo que o anterior, adiciona o texto ao bloco atual
        if speaker == current_speaker:
            current_text += " " + text
        else:
            # Se o falante é diferente, salva o bloco anterior (se existir) e inicia um novo
            if current_speaker is not None:
                grouped_transcription.append(
                    {"speaker": current_speaker, "text": current_text.strip()}
                )

            # Atualiza o falante e o texto atual
            current_speaker = speaker
            current_text = text

    # Adiciona o último bloco de texto, se existir
    if current_speaker is not None:
        grouped_transcription.append(
            {"speaker": current_speaker, "text": current_text.strip()}
        )

    return grouped_transcription


def save_transcription_to_db(file_name, transcription_text, model, execution_time):
    # Executa a query com os dados passados
    with conn.session as session:
        session.execute(
            text(
                "INSERT INTO transcriptions (file_name, transcription, model, execution_time) VALUES(:file_name, :transcription, :model, :execution_time);"
            ),
            {
                "file_name": file_name,
                "transcription": transcription_text,
                "model": model,
                "execution_time": execution_time,
            },
        )
        session.commit()  # Confirma a transação


def upload_view():
    st.title("Realize uma nova Transcrição")

    st.markdown(
        "Clique em `Browse files` para buscar no computador o áudio desejado. Você também tem a opção de arrastar e soltar o arquivo para a área de upload."
    )

    st.markdown(
        "Os modelos de transcrição vão do ``tiny`` ao ``large``. Quanto maior a precisão/confiabilidade (mais próximo de ``large``), mais tempo será necessário para processar sua solicitação."
    )

    speaker_colors = {
        "SPEAKER 1": "FF0000",  # Vermelho
        "SPEAKER 2": "00FF00",  # Verde
        "SPEAKER 3": "0000FF",  # Azul
        "SPEAKER 4": "FFFF00",  # Amarelo
        "SPEAKER 5": "FF00FF",  # Magenta
        "SPEAKER 6": "00FFFF",  # Ciano
        "SPEAKER 7": "FFA500",  # Laranja
        "SPEAKER 8": "800080",  # Roxo
        "SPEAKER 9": "808080",  # Cinza
        "SPEAKER 10": "000000",  # Preto
    }

    # Formulário para upload de arquivo e seleção de modelo
    with st.form("input_form"):
        input_file = st.file_uploader(
            "Arquivos de áudio",
            type=["mp4", "m4a", "mp3", "wav"],
            accept_multiple_files=False,
        )

        # Permite selecionar o modelo Whisper (de "tiny" a "large").

        whisper_model = st.selectbox(
            "Modelo de Transcriçao",
            options=["tiny", "base", "small", "medium", "large"],
            index=4,
        )

        num_speakers = st.number_input(
            "Quantidade de falantes", min_value=1, max_value=10
        )

        btn_transcribe = st.form_submit_button(label="Iniciar")

    if btn_transcribe:
        # Se o usuário clicar em "Iniciar" e houver arquivos carregados, a transcrição será inicializada
        if input_file:
            start_time = time.time()

            # print(input_file.name)
            with st.spinner("Transcrevendo o áudio..."):
                segments = transcribe(input_file, whisper_model, num_speakers)

            end_time = time.time()
            execution_time = end_time - start_time
            # print(execution_time)

            st.success("Transcrição finalizada!")
            st.write(f"Tempo de execução: {execution_time:.2f} segundos")

            grouped_segments = group_speaker_segments(segments)

            st.markdown("### Transcrição:")

            text_content = ""
            file_name = (
                input_file.name
                + "-"
                + whisper_model
                + "-"
                + datetime.today().strftime("%d-%m-%y")
                + ".docx"
            )

            for segment in grouped_segments:
                speaker = segment["speaker"]
                text = segment["text"]
                color = speaker_colors.get(speaker, "000000")  # Cor padrão é preto

                # Exibe o texto colorido
                st.markdown(
                    f"<span><strong style='color: #{color};'>{speaker}:</strong> {text}</span>",
                    unsafe_allow_html=True,
                )

                text_content += f"{speaker}: {text}\n"

            existing_transcription = fetch_transcription_by_file_name(file_name)

            if existing_transcription is None:
                save_transcription_to_db(
                    file_name, text_content, whisper_model, execution_time
                )

            # Cria um arquivo docx na memória
            doc = docx.Document()

            # Adiciona os segmentos ao documento com cores
            for segment in segments:
                speaker = segment["speaker"]
                text = segment["text"]
                color = speaker_colors.get(speaker, "000000")  # Cor padrão é preto

                # Adiciona um parágrafo com a cor do texto
                paragraph = doc.add_paragraph()

                # Adiciona o nome do falante em cor
                speaker_run = paragraph.add_run(f"{speaker}: ")
                speaker_run.font.color.rgb = docx.shared.RGBColor(
                    int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16)
                )

                # Adiciona o texto do falante em preto
                text_run = paragraph.add_run(text)
                text_run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)  # Preto

            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)

            # Botão para baixar o arquivo transcrito
            st.download_button(
                label="Baixar Transcrição",
                data=bio.getvalue(),
                file_name=file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        else:
            st.error("Por favor, selecione um arquivo.")


def main():

    _style_language_uploader()

    st.sidebar.subheader("Navegação no sistema")

    st.sidebar.markdown(
        "Clique em `Nova transcrição` para realizar transcrição de um novo arquivo. Caso deseje consultar transcrições já realizadas, clique em `Histórico`."
    )

    st.sidebar.subheader("Download de transcrições")
    st.sidebar.markdown(
        "Clique em ``Baixar Transcrição`` para obter o arquivo em .docx."
    )

    pg = st.navigation(
        [
            st.Page(
                upload_view,
                title="Nova transcrição",
                icon=":material/insert_drive_file:",
            ),
            st.Page(
                display_transcriptions, title="Histórico", icon=":material/history:"
            ),
        ]
    )
    pg.run()


if __name__ == "__main__":
    main()
